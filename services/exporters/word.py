from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_word(search_data, schools, map_image_bytes=None):
    """워드 문서 생성"""
    doc = Document()

    # 한글 폰트 지정 (Windows 기본: 맑은 고딕)
    font_name = 'Malgun Gothic'
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def set_run_font(run):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

    def apply_font_to_paragraph(paragraph):
        for run in paragraph.runs:
            set_run_font(run)

    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run('업 무 보 고')
    title_run.bold = True
    title_run.font.size = Pt(20)
    set_run_font(title_run)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 부제목
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('제목: 재난 발생에 따른 인근 학교 현황')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    set_run_font(subtitle_run)

    doc.add_paragraph()

    # 1. 개요
    section1 = doc.add_paragraph()
    section1_run = section1.add_run('1. 개요')
    section1_run.bold = True
    section1_run.font.size = Pt(12)
    set_run_font(section1_run)

    doc.add_paragraph(f"    - 일시: {search_data['datetime']}")
    doc.add_paragraph(f"    - 장소: {search_data['address']}")
    doc.add_paragraph(f"    - 좌표: (위도 {search_data['lat']}, 경도 {search_data['lng']})")

    doc.add_paragraph()

    # 2. 인근 학교 현황
    section2 = doc.add_paragraph()
    section2_run = section2.add_run('2. 인근 학교 현황')
    section2_run.bold = True
    section2_run.font.size = Pt(12)
    set_run_font(section2_run)

    # 학교급별 집계
    elem = sum(1 for s in schools if '초등' in s.get('type', ''))
    middle = sum(1 for s in schools if '중학' in s.get('type', ''))
    high = sum(1 for s in schools if '고등' in s.get('type', ''))
    special = sum(1 for s in schools if '특수' in s.get('type', ''))

    doc.add_paragraph(f"    - 반경 {search_data['radius']}km 이내 총 {len(schools)}개교")

    # 집계 문자열 생성 (0개인 학교급은 제외)
    counts = []
    if elem > 0:
        counts.append(f"초등학교 {elem}개교")
    if middle > 0:
        counts.append(f"중학교 {middle}개교")
    if high > 0:
        counts.append(f"고등학교 {high}개교")
    if special > 0:
        counts.append(f"특수학교 {special}개교")

    if counts:
        doc.add_paragraph(f"      ({', '.join(counts)})")

    doc.add_paragraph()

    # 표 생성
    if schools:
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'

        # 헤더
        header_cells = table.rows[0].cells
        headers = ['순번', '학교명', '학교급', '학급수', '학생수', '거리(km)', '도로명주소', '비고']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    set_run_font(run)

        # 데이터 행
        for idx, school in enumerate(schools, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = school['name']
            row[2].text = school['type']
            row[3].text = str(school.get('classes', 0))
            row[4].text = str(school.get('students', 0))
            row[5].text = str(school['distance'])
            row[6].text = school.get('road_address', '') or school.get('address', '')
            row[7].text = ''

    doc.add_paragraph()

    # 3. 조치사항
    section3 = doc.add_paragraph()
    section3_run = section3.add_run('3. 조치사항')
    section3_run.bold = True
    section3_run.font.size = Pt(12)
    set_run_font(section3_run)

    doc.add_paragraph("    - ")
    doc.add_paragraph("    - ")
    doc.add_paragraph("    - ")

    doc.add_paragraph()

    # 4. 붙임 (지도 이미지)
    if map_image_bytes:
        section4 = doc.add_paragraph()
        section4_run = section4.add_run('4. 붙임: 위치도 1부')
        section4_run.bold = True
        section4_run.font.size = Pt(12)
        set_run_font(section4_run)

        doc.add_paragraph()
        doc.add_picture(map_image_bytes, width=Inches(6))

    for paragraph in doc.paragraphs:
        apply_font_to_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_font_to_paragraph(paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
